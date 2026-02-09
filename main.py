from dotenv import load_dotenv
load_dotenv()

import customtkinter as ctk
from tkinter import messagebox
from groq import Groq
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import speech_recognition as sr
import threading
import time
import winsound
import os
import sqlite3
from datetime import datetime

# =====================
# CONFIG
# =====================

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# =====================
# BASE DE DATOS HISTORIAL
# =====================

conn = sqlite3.connect("historial.db")
cursor = conn.cursor()

cursor.execute("""
CREATE TABLE IF NOT EXISTS historial(
    id INTEGER PRIMARY KEY,
    fecha TEXT,
    tipo TEXT,
    contenido TEXT
)
""")

conn.commit()

# =====================
# CONFIG VOZ
# =====================

recognizer = sr.Recognizer()
grabando = False
parpadeo_activo = False

# =====================
# SONIDOS
# =====================

def beep_inicio():
    winsound.Beep(1000,200)

def beep_fin():
    winsound.Beep(600,300)

# =====================
# PARPADEO
# =====================

def parpadear():
    global parpadeo_activo
    if grabando:
        indicador.configure(text_color="red" if parpadeo_activo else "gray")
        parpadeo_activo = not parpadeo_activo
        app.after(500, parpadear)

# =====================
# DICTADO
# =====================

def iniciar_dictado():
    global grabando
    grabando = True
    beep_inicio()
    indicador.configure(text="Grabando...")
    parpadear()
    threading.Thread(target=dictado_continuo).start()

def detener_dictado():
    global grabando
    grabando = False
    beep_fin()
    indicador.configure(text="Dictado detenido", text_color="gray")

def dictado_continuo():
    global grabando
    with sr.Microphone() as source:
        inicio = time.time()

        while grabando:
            try:
                audio = recognizer.listen(source, phrase_time_limit=5)
                texto = recognizer.recognize_google(audio, language="es-CO")
                textbox.insert("end", texto+" ")

                if time.time()-inicio > 10:
                    if not messagebox.askyesno("Dictado","¿Continuar dictando?"):
                        detener_dictado()
                        break
                    inicio = time.time()
            except:
                pass

# =====================
# IA
# =====================

def generar_texto():

    instrucciones = textbox.get("1.0","end").strip()
    tipo = tipo_documento.get()

    if instrucciones == "":
        return

    prompt = f"Crear contenido educativo tipo {tipo} sobre: {instrucciones}"

    r = client.chat.completions.create(
        messages=[{"role":"user","content":prompt}],
        model="llama-3.3-70b-versatile"
    )

    texto = r.choices[0].message.content

    resultado.delete("1.0","end")
    resultado.insert("1.0",texto)

    # Guardar historial
    cursor.execute(
        "INSERT INTO historial VALUES(NULL,?,?,?)",
        (datetime.now().strftime("%d/%m/%Y"),tipo,texto)
    )
    conn.commit()

# =====================
# DOCUMENTO
# =====================

def crear_documento(tipo, contenido):

    doc = Document()

    tabla = doc.add_table(rows=1, cols=2)

    try:
        tabla.cell(0,0).paragraphs[0].add_run().add_picture("logo.png", width=Inches(1.1))
    except:
        pass

    header = tabla.cell(0,1).paragraphs[0]
    header.add_run("INSTITUTO LUIS CARLOS GALÁN SARMIENTO\n").bold=True
    header.add_run(f"{tipo}\n")
    header.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_paragraph(contenido)

    return doc

# =====================
# GUARDAR
# =====================

def guardar_word():
    doc = crear_documento(tipo_documento.get(), resultado.get("1.0","end"))
    doc.save("Documento.docx")
    messagebox.showinfo("Guardado","Word creado")

def guardar_pdf():
    doc = crear_documento(tipo_documento.get(), resultado.get("1.0","end"))
    doc.save("temp.docx")
    convert("temp.docx","Documento.pdf")
    os.remove("temp.docx")
    messagebox.showinfo("Guardado","PDF creado")

# =====================
# VISTA PREVIA
# =====================

def vista_previa():

    ventana = ctk.CTkToplevel(app)
    ventana.geometry("600x400")
    ventana.title("Vista previa")

    texto = resultado.get("1.0","end")

    preview = ctk.CTkTextbox(ventana,width=550,height=350)
    preview.pack(pady=20)
    preview.insert("1.0", texto)

# =====================
# HISTORIAL
# =====================

def abrir_historial():

    ventana = ctk.CTkToplevel(app)
    ventana.geometry("600x400")
    ventana.title("Historial")

    lista = ctk.CTkTextbox(ventana,width=550,height=350)
    lista.pack(pady=20)

    cursor.execute("SELECT fecha,tipo FROM historial")
    registros = cursor.fetchall()

    for r in registros:
        lista.insert("end", f"{r[0]} - {r[1]}\n")

# =====================
# DASHBOARD
# =====================

def dashboard():

    ventana = ctk.CTkToplevel(app)
    ventana.geometry("400x250")
    ventana.title("Dashboard")

    cursor.execute("SELECT COUNT(*) FROM historial")
    total = cursor.fetchone()[0]

    ctk.CTkLabel(ventana,text="Bienvenido al sistema",
                 font=("Arial",18,"bold")).pack(pady=20)

    ctk.CTkLabel(ventana,text=f"Documentos generados: {total}",
                 font=("Arial",16)).pack()

# =====================
# INTERFAZ
# =====================

app = ctk.CTk()
app.geometry("1000x650")
app.title("Sistema Generador Académico IA")

sidebar = ctk.CTkFrame(app,width=250)
sidebar.pack(side="left",fill="y")

ctk.CTkLabel(sidebar,text="Panel IA",font=("Arial",20,"bold")).pack(pady=20)

tipo_documento = ctk.StringVar(value="Guía de Aprendizaje")

ctk.CTkOptionMenu(sidebar,
                  values=["Guía de Aprendizaje","Ficha de Ejercicios"],
                  variable=tipo_documento).pack(pady=10)

indicador = ctk.CTkLabel(sidebar,text="Dictado detenido")
indicador.pack(pady=10)

ctk.CTkButton(sidebar,text="🎤 Dictar",command=iniciar_dictado).pack(pady=5)
ctk.CTkButton(sidebar,text="🛑 Detener",command=detener_dictado).pack(pady=5)

ctk.CTkButton(sidebar,text="Vista previa",command=vista_previa).pack(pady=15)
ctk.CTkButton(sidebar,text="Historial",command=abrir_historial).pack()
ctk.CTkButton(sidebar,text="Dashboard",command=dashboard).pack(pady=15)

ctk.CTkButton(sidebar,text="Guardar Word",command=guardar_word).pack(pady=5)
ctk.CTkButton(sidebar,text="Guardar PDF",command=guardar_pdf).pack(pady=5)

main = ctk.CTkFrame(app)
main.pack(side="right",expand=True,fill="both")

ctk.CTkLabel(main,text="Instrucciones",font=("Arial",18,"bold")).pack(pady=10)

textbox = ctk.CTkTextbox(main,width=600,height=150)
textbox.pack()

ctk.CTkButton(main,text="Generar IA",command=generar_texto,
              height=40,font=("Arial",14,"bold")).pack(pady=10)

resultado = ctk.CTkTextbox(main,width=600,height=250)
resultado.pack()

app.mainloop()






